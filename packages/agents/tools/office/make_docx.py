#!/usr/bin/env python3
"""Build a Word document from markdown, or from a JSON spec. One command:

    make-doc brief.md out.docx

    # Project Brief          <- the document title
    ## Background            <- a section heading
    One paragraph.           <- a paragraph
    - A bullet.              <- a bullet

### is a sub-heading. Everything else is a paragraph. Nothing needs quoting
or escaping, which is the point: the JSON form below is still accepted, and
a model writing a long paragraph into it dropped the closing quote and then
repeated the same broken spec seven times. Prose does not survive being
hand-typed into JSON; markdown has nothing to get wrong.

    {"title": "Project Brief",
     "blocks": [
       {"style": "h1", "text": "Background"},
       {"style": "p",  "text": "One paragraph."},
       {"style": "li", "text": "A bullet."}
     ]}

Styles: h1 h2 p li. An unknown style stops the build with the list of valid
ones, because a block silently dropped becomes a document that looks
finished and is missing a section.

This exists because a model writing python-docx code from memory invents
imports and APIs; a model writing markdown and running one command does not.
"""
from __future__ import annotations

import json
import os
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from artifact_path import resolve_output  # noqa: E402

STYLES = ("h1", "h2", "p", "li")


def build(spec: dict, out: Path) -> int:
    from docx import Document

    doc = Document()
    title = spec.get("title")
    if title:
        doc.add_heading(str(title), level=0)
    blocks = spec.get("blocks")
    if not isinstance(blocks, list) or not blocks:
        raise SystemExit('spec has no "blocks": expected a non-empty list of {"style","text"}')
    for i, block in enumerate(blocks):
        style = block.get("style") if isinstance(block, dict) else None
        text = str(block.get("text", "")) if isinstance(block, dict) else ""
        if style == "h1":
            doc.add_heading(text, level=1)
        elif style == "h2":
            doc.add_heading(text, level=2)
        elif style == "p":
            doc.add_paragraph(text)
        elif style == "li":
            doc.add_paragraph(text, style="List Bullet")
        else:
            raise SystemExit(f'blocks[{i}] has style {style!r}: use one of {", ".join(STYLES)}')
    parent = os.path.dirname(str(out))
    if parent:
        os.makedirs(parent, exist_ok=True)
    doc.save(str(out))
    print(f"wrote {out} ({len(blocks)} blocks)")
    return 0


def from_markdown(text: str) -> dict:
    """A spec out of the markdown subset above.

    Deliberately forgiving in one direction only: anything it does not
    recognise becomes a paragraph, so no line is ever dropped. A dropped line
    is the failure mode that matters — a document that looks complete and is
    missing a sentence.
    """
    spec: dict = {"blocks": []}
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            spec["blocks"].append({"style": "h2", "text": line[4:].strip()})
        elif line.startswith("## "):
            spec["blocks"].append({"style": "h1", "text": line[3:].strip()})
        elif line.startswith("# "):
            head = line[2:].strip()
            # The first # is the document's title; a second one is a section,
            # because a document with two titles is not a thing.
            if "title" not in spec:
                spec["title"] = head
            else:
                spec["blocks"].append({"style": "h1", "text": head})
        elif line.startswith(("- ", "* ")):
            spec["blocks"].append({"style": "li", "text": line[2:].strip()})
        else:
            spec["blocks"].append({"style": "p", "text": line})
    return spec


def main(argv: list[str]) -> int:
    # A wrong count gets a pointed sentence, not the usage dump: the one way
    # this really happens is an unquoted space in the output path, and a model
    # shown the whole docstring retried the same command a dozen times where a
    # model told "quote it or drop the space" fixed it in one.
    if len(argv) > 3:
        raise SystemExit(
            f"got {len(argv) - 1} arguments, expected 2: make-doc spec.json out.docx —"
            " a path with a space must be quoted; better, name the file with dashes,"
            " like /artifacts/docs/project-brief.docx"
        )
    if len(argv) < 3:
        print(__doc__)
        return 2
    spec_path, out = Path(argv[1]), Path(resolve_output(argv[2]))
    try:
        text = spec_path.read_text("utf8")
    except FileNotFoundError:
        raise SystemExit(f"{spec_path}: no such file — write the source first, then run make-doc")
    # Which format by what is in the file, never by the extension: a model
    # that writes markdown into spec.json meant the markdown.
    if text.lstrip().startswith("{"):
        try:
            spec = json.loads(text)
        except json.JSONDecodeError as e:
            raise SystemExit(
                f"{spec_path} starts with {{ so it was read as JSON, and it is not valid: {e}."
                " Write it as markdown instead — # title, ## section, - bullet, anything"
                " else a paragraph — which needs no quoting and cannot fail this way."
            )
        if not isinstance(spec, dict):
            raise SystemExit(f"{spec_path}: expected a JSON object, got {type(spec).__name__}")
    else:
        spec = from_markdown(text)
    return build(spec, out)


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
