# The make-doc builder. One function, one record argument — the spec — and
# the path it writes is in the spec, under /artifacts, where the run's
# reconcile finds it.
#
# spec = {
#   "path": "/artifacts/docs/plan.docx",   # required
#   "title": "Q3 Plan",                    # document heading
#   "subtitle": "",                        # optional line under the title
#   "sections": [                          # in order
#     { "heading": "Context", "paragraphs": ["...", "..."],
#       "bullets": ["...", "..."] },       # paragraphs then bullets, either optional
#   ],
# }
import os
from docx import Document
from docx.shared import Pt

def build(spec):
    doc = Document()
    doc.add_heading(spec.get("title", "Untitled"), level=0)
    if spec.get("subtitle"):
        p = doc.add_paragraph(spec["subtitle"])
        p.runs[0].font.size = Pt(12)
    for sect in spec.get("sections", []):
        if sect.get("heading"):
            doc.add_heading(sect["heading"], level=1)
        for text in sect.get("paragraphs", []):
            doc.add_paragraph(text)
        for item in sect.get("bullets", []):
            doc.add_paragraph(item, style="List Bullet")
    path = spec["path"]
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    print("wrote", path)
