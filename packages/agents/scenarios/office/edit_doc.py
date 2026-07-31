# Fill a document that already exists.
#
# The counterpart to build_doc: build makes a file from nothing, this one
# opens the template the conversation was started from and replaces its
# <PLACEHOLDER> runs in place. That difference is the whole point of a
# template — its styles, its heading levels, its table borders are already
# decided, and a document generated beside it throws all of that away.
#
# Saving back to the same path is what makes the run's reconcile append
# version 2 of the same artifact rather than create a second one, so the
# reader sees the document they picked, filled in, with its history.
#
# fills = { "<TITLE>": "Q3 status", "<PERIOD>": "July–September", ... }
# Anything left unfilled stays visibly angled rather than becoming an empty
# line — an unanswered placeholder should look unanswered.
import os
from docx import Document


def _replace_in_paragraph(paragraph, fills):
    # Runs split mid-word, so a placeholder can straddle several of them.
    # Joining the paragraph and rewriting the first run keeps the paragraph's
    # style (its heading level, its list bullet) while replacing the text.
    text = "".join(run.text for run in paragraph.runs)
    if not text:
        return False
    out = text
    for key, value in fills.items():
        out = out.replace(key, str(value))
    if out == text:
        return False
    for run in paragraph.runs[1:]:
        run.text = ""
    paragraph.runs[0].text = out
    return True


def fill(path, fills, extra_bullets=None):
    """Replace placeholders in `path`, in place.

    extra_bullets = { "What happened": ["...", "..."] } appends bullets under
    a heading, for the case where the template ships three <item> lines and
    the truth needs six. The heading is matched by its text, so a template
    whose headings were renamed needs the new names here.
    """
    doc = Document(path)
    hits = 0
    for paragraph in doc.paragraphs:
        if _replace_in_paragraph(paragraph, fills):
            hits += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if _replace_in_paragraph(paragraph, fills):
                        hits += 1

    for heading, items in (extra_bullets or {}).items():
        anchor = None
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip() == heading:
                anchor = i
        if anchor is None:
            continue
        # Insert after the last existing bullet under that heading, so added
        # lines land with the template's own bullets and not at the end of
        # the document.
        last = anchor
        for i in range(anchor + 1, len(doc.paragraphs)):
            if doc.paragraphs[i].style.name.startswith("Heading"):
                break
            last = i
        for item in items:
            new = doc.add_paragraph(item, style="List Bullet")
            doc.paragraphs[last]._p.addnext(new._p)
            last += 1

    doc.save(path)
    print(f"filled {hits} placeholders in {os.path.basename(path)}")
    return hits
